"""
Script to create a Word document version of the USER_TUTORIAL.md file.
"""
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re

logger = logging.getLogger(__name__)


class TutorialWordGenerator:
    """Generate Word document from markdown tutorial."""
    
    def __init__(self):
        """Initialize the Word generator."""
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Setup custom styles for the document."""
        # Title style
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = None  # Default color
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1 style
        h1_style = self.doc.styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Arial'
        h1_font.size = Pt(18)
        h1_font.bold = True
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
        
        # Heading 2 style
        h2_style = self.doc.styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Arial'
        h2_font.size = Pt(16)
        h2_font.bold = True
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(4)
        
        # Heading 3 style
        h3_style = self.doc.styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Arial'
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_style.paragraph_format.space_before = Pt(8)
        h3_style.paragraph_format.space_after = Pt(4)
        
        # Code style
        code_style = self.doc.styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Consolas'
        code_font.size = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
    
    def parse_markdown_to_word(self, markdown_file: Path, output_file: Path):
        """Convert markdown file to Word document."""
        try:
            with open(markdown_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split content into lines
            lines = content.split('\n')
            
            # Add title page
            self.add_title_page()
            
            # Process each line
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                if not line:
                    i += 1
                    continue
                
                # Handle headers
                if line.startswith('# '):
                    self.add_heading(line[2:], 1)
                elif line.startswith('## '):
                    self.add_heading(line[3:], 2)
                elif line.startswith('### '):
                    self.add_heading(line[4:], 3)
                elif line.startswith('#### '):
                    self.add_heading(line[5:], 4)
                elif line.startswith('##### '):
                    self.add_heading(line[6:], 5)
                
                # Handle code blocks
                elif line.startswith('```'):
                    i = self.handle_code_block(lines, i)
                    continue
                
                # Handle lists
                elif line.startswith('- ') or line.startswith('* '):
                    i = self.handle_list(lines, i, ordered=False)
                    continue
                elif re.match(r'^\d+\. ', line):
                    i = self.handle_list(lines, i, ordered=True)
                    continue
                
                # Handle regular paragraphs
                else:
                    if line and not line.startswith('#'):
                        self.add_paragraph(line)
                
                i += 1
            
            # Save document
            self.doc.save(output_file)
            logger.info(f"Word document created: {output_file}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            return False
    
    def add_title_page(self):
        """Add a professional title page."""
        # Main title
        title = self.doc.add_paragraph()
        title.style = 'CustomTitle'
        title_run = title.runs[0] if title.runs else title.add_run()
        title_run.text = "Automation Studio Selector"
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Complete User Tutorial")
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.italic = True
        
        # Version info
        self.doc.add_paragraph()  # Empty line
        version = self.doc.add_paragraph()
        version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_run = version.add_run("Version 1.1.0")
        version_run.font.size = Pt(14)
        
        # Author info
        self.doc.add_paragraph()  # Empty line
        author = self.doc.add_paragraph()
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author.add_run("Created by: Vitaly Grosman\nIndigo R&D Division\n© 2025")
        author_run.font.size = Pt(12)
        
        # Page break
        self.doc.add_page_break()
    
    def add_heading(self, text, level):
        """Add a heading with appropriate style."""
        # Clean text (remove emojis and markdown formatting)
        clean_text = re.sub(r'[🎯🚀🖥️⚡🔧💡📋📁📝📊🏠⚙️🔄👁️📦🔍🛡️⏰🔴🚪📄💾🔘❌✅📚🎨📐🎛️🔧]', '', text).strip()
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)  # Remove bold markdown
        
        heading = self.doc.add_paragraph()
        if level == 1:
            heading.style = 'CustomH1'
        elif level == 2:
            heading.style = 'CustomH2'
        elif level == 3:
            heading.style = 'CustomH3'
        else:
            heading.style = 'CustomH3'
        
        heading_run = heading.add_run(clean_text)
    
    def add_paragraph(self, text):
        """Add a regular paragraph."""
        if not text.strip():
            return
        
        # Clean text
        clean_text = self.clean_markdown_text(text)
        
        if clean_text.strip():
            para = self.doc.add_paragraph()
            para.add_run(clean_text)
            para.paragraph_format.space_after = Pt(6)
    
    def handle_code_block(self, lines, start_index):
        """Handle code block formatting."""
        i = start_index + 1
        code_lines = []
        
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        if code_lines:
            code_para = self.doc.add_paragraph()
            code_para.style = 'CustomCode'
            code_para.add_run('\n'.join(code_lines))
        
        return i + 1
    
    def handle_list(self, lines, start_index, ordered=False):
        """Handle list formatting."""
        i = start_index
        list_items = []
        
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            if (not ordered and (line.startswith('- ') or line.startswith('* '))) or \
               (ordered and re.match(r'^\d+\. ', line)):
                # Extract list item text
                if ordered:
                    item_text = re.sub(r'^\d+\. ', '', line)
                else:
                    item_text = line[2:]  # Remove '- ' or '* '
                
                list_items.append(self.clean_markdown_text(item_text))
                i += 1
            else:
                break
        
        # Add list to document
        for item in list_items:
            if item.strip():
                para = self.doc.add_paragraph()
                para.style = 'List Bullet' if not ordered else 'List Number'
                para.add_run(item)
        
        return i
    
    def clean_markdown_text(self, text):
        """Clean markdown formatting from text."""
        # Remove emojis
        text = re.sub(r'[🎯🚀🖥️⚡🔧💡📋📁📝📊🏠⚙️🔄👁️📦🔍🛡️⏰🔴🚪📄💾🔘❌✅📚🎨📐🎛️🔧]', '', text)
        
        # Convert bold markdown to regular text (Word will handle formatting)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        
        # Convert italic markdown
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        
        # Convert inline code
        text = re.sub(r'`(.*?)`', r'\1', text)
        
        # Clean up extra spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()


def create_word_tutorial():
    """Main function to create the Word tutorial."""
    try:
        # Check if python-docx is available
        try:
            from docx import Document
        except ImportError:
            print("ERROR: python-docx library is required.")
            print("Install it with: pip install python-docx")
            return False
        
        generator = TutorialWordGenerator()
        
        markdown_file = Path("USER_TUTORIAL.md")
        word_file = Path("Automation_Studio_Selector_Tutorial.docx")
        
        if not markdown_file.exists():
            print(f"ERROR: {markdown_file} not found")
            return False
        
        print("Creating Word document from USER_TUTORIAL.md...")
        success = generator.parse_markdown_to_word(markdown_file, word_file)
        
        if success:
            print(f"SUCCESS: Word document created as {word_file}")
            print(f"File location: {word_file.absolute()}")
            return True
        else:
            print("ERROR: Failed to create Word document")
            return False
            
    except Exception as e:
        print(f"ERROR: {e}")
        return False


if __name__ == "__main__":
    create_word_tutorial()
